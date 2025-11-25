from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict


class DocxService:
    """Service for generating Word documents"""
    
    def generate_document(self, title: str, sections: List[Dict[str, str]]) -> BytesIO:
        """
        Generate a .docx file from sections
        
        Args:
            title: Document title
            sections: List of dicts with 'title' and 'content' keys
            
        Returns:
            BytesIO object containing the document
        """
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sections
        for section in sections:
            # Add section heading
            doc.add_heading(section['title'], level=1)
            
            # Add content
            if section.get('content'):
                # Split content into paragraphs
                paragraphs = section['content'].split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.line_spacing = 1.15
        
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream


# Singleton instance
docx_service = DocxService()
